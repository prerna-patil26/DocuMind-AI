from datetime import datetime
import streamlit as st
from PyPDF2 import PdfReader
import os
import pytesseract
from PIL import Image
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate
from langchain.chains.question_answering import load_qa_chain
from dotenv import load_dotenv
import io
from fpdf import FPDF
import docx
import re
import base64
import fitz  # PyMuPDF
import shutil
import glob
from streamlit.components.v1 import html

# Configure Tesseract path
current_dir = os.path.dirname(os.path.abspath(__file__))
pytesseract.pytesseract.tesseract_cmd = os.path.join(current_dir, "tesseract", "tesseract.exe")

# Load environment variables
load_dotenv()

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'processed_text' not in st.session_state:
    st.session_state.processed_text = ""
if 'show_guide' not in st.session_state:
    st.session_state.show_guide = False
if 'download_format' not in st.session_state:
    st.session_state.download_format = "TXT"
if 'original_files' not in st.session_state:
    st.session_state.original_files = None
if 'show_project_info' not in st.session_state:
    st.session_state.show_project_info = False

# Custom CSS for styling
st.markdown("""
<style>
/* 1. Remove button blinking and add glow effect */
button, .stButton>button {
    background: linear-gradient(to right, #7b1fa2, #2196f3) !important;
    color: white !important;
    border: none !important;
    border-radius: 25px !important;
    padding: 6px 16px !important;
    font-size: 0.85rem !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 0 5px rgba(123, 31, 162, 0.3);
    animation: button-glow 2s infinite alternate;
}

@keyframes button-glow {
    0% { box-shadow: 0 0 5px rgba(123, 31, 162, 0.3); }
    100% { box-shadow: 0 0 15px rgba(33, 150, 243, 0.5); }
}

button:hover, .stButton>button:hover {
    opacity: 0.9;
    transform: scale(1.02);
}

/* 2. Sidebar section styling */
.sidebar-section-title {
    text-align: center;
    background: linear-gradient(to right, #7f5af0, #3c8ce7);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-weight: 600;
    margin-bottom: 0.5rem;
}

.sidebar-section {
    border-radius: 15px !important;
    padding: 1rem !important;
    margin-bottom: 1rem !important;
    background: rgba(30, 30, 30, 0.7) !important;
}

.sidebar-divider {
    height: 2px;
    background: linear-gradient(to right, #7b1fa2, #2196f3);
    margin: 1rem 0;
    border: none;
}

.sidebar-header {
    font-size: 1.8rem;
    font-weight: 700;
    background: linear-gradient(to right, #7f5af0, #3c8ce7);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    text-align: center;
    margin-bottom: 0.5rem;
    text-shadow: 0 0 8px rgba(123, 31, 162, 0.4);
    animation: header-glow 2s infinite alternate;
}
    
[data-testid="stSidebar"] {
    position: relative;
}

/* Create the glowing gradient border */
[data-testid="stSidebar"]::after {
    content: "";
    position: absolute;
    right: 0;
    top: 0;
    height: 100%;
    width: 3px;
    background: linear-gradient(to bottom, 
        #7b1fa2, 
        #9c27b0, 
        #2196f3);
    box-shadow: 
        0 0 5px #7b1fa2,
        0 0 10px #2196f3;
    animation: border-glow 2s ease-in-out infinite alternate;
}

@keyframes border-glow {
    from { opacity: 0.8; }
    to { opacity: 1; }
}

/* Ensure main content doesn't overlap */
.main .block-container {
    padding-left: 15px;
}

.sidebar-subheader {
    font-size: 0.9rem;
    color: #aaa;
    text-align: center;
    margin-bottom: 1.5rem;
}

[data-testid="stSidebar"] {
    background-color: #000000 !important;
    width: 350px !important;
}
    
.info-card {
    padding: 1rem;
    margin-bottom: 1.5rem;
    background: rgba(30, 30, 30, 0.5);
    border-left: 4px solid;
    border-image: linear-gradient(to bottom, #7b1fa2, #2196f3) 1;
    animation: card-glow 3s infinite;
    border-radius: 15px !important;
    overflow: hidden !important;
}

@keyframes card-glow {
    0% { box-shadow: 0 0 5px rgba(123, 31, 162, 0.3); }
    50% { box-shadow: 0 0 15px rgba(33, 150, 243, 0.5); }
    100% { box-shadow: 0 0 5px rgba(123, 31, 162, 0.3); }
}

.info-card h3 {
    margin-top: 0;
    background: linear-gradient(to right, #7f5af0, #3c8ce7);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

/* 3. File uploader styling */
div[data-testid="stFileUploader"] {
    border: 2px solid;
    border-image: linear-gradient(to right, #7b1fa2, #2196f3);
    border-image-slice: 1;
    border-radius: 10px;
    padding: 10px;
    border-radius: 15px !important;
    overflow: hidden !important;
    padding: 10px !important;
    background: rgba(30, 30, 30, 0.7) !important;

}

div[data-testid="stFileUploader"] button {
    margin: 0 auto !important;
    display: block !important;
    border-radius: 20px !important;
}

/* Style the file list items */
div[data-testid="stFileUploader"] [role='listitem'] {
    border-radius: 8px !important;
    margin: 0.25rem 0 !important;
}
            
div[data-testid="stFileUploader"] button[title="Delete"] {
    background: transparent !important;
    color: #ff4b4b !important;
    border: none !important;
    box-shadow: none !important;
    padding: 0 !important;
    margin: 0 !important;
}

div[data-testid="stFileUploader"] > div:first-child {
    background: linear-gradient(to right, #7b1fa2, #2196f3) !important;
    border-radius: 20px !important;
    padding: 1rem 1.2rem !important;
    color: white !important;
    font-weight: 600 !important;
    border: none !important;
    text-align: center !important;
    margin: 0 auto !important;
    max-width: 90% !important;
    border-radius: 15px !important;
}
/* 4. Toast notification styling */
.toast {
    font-size: 0.9rem;
    padding: 0.75rem;
    border-radius: 8px;
}

/* Chat message styling */
.user-message-container {
    display: flex;
    justify-content: flex-end;
    margin-bottom: 1rem;
}

.bot-message-container {
    display: flex;
    justify-content: flex-start;
    margin-bottom: 1rem;
}

.user-message {
    max-width: 80%;
    padding: 0.75rem 1rem;
    border-radius: 12px;
    background: linear-gradient(to left, rgba(60, 140, 231, 0.1), transparent);
    border-left: 4px solid #3c8ce7;
    border-right: 4px solid #3c8ce7;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}

.bot-message {
    max-width: 80%;
    padding: 0.75rem 1rem;
    border-radius: 12px;
    background: linear-gradient(to right, rgba(127, 90, 240, 0.1), transparent);
    border-left: 4px solid #7f5af0;
    border-right: 4px solid #7f5af0;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}

.message-meta {
    font-size: 0.75rem;
    color: #666;
    margin-bottom: 0.25rem;
}

.message-content {
    word-wrap: break-word;
}

/* 6. Download dropdown styling */
.download-container {
    display: flex;
    align-items: center;
    gap: 0.5rem;
}

.download-label {
    font-size: 0.9rem;
    margin-bottom: 0;
}

.download-select {
    flex-grow: 1;
}

.download-button {
    min-width: 40px;
}

/* 7. Developer info styling */
.developer-info {
    text-align: center;
    margin-top: 1.5rem;
    padding-top: 1rem;
    border-top: 1px solid #444;
}

.developer-line {
    display: flex;
    justify-content: center;
    align-items: center;
    gap: 1rem;
    margin-bottom: 1rem;
    flex-wrap: wrap;
}
            
.developer-name {
    font-weight: 600;
    background: linear-gradient(to right, #7f5af0, #3c8ce7);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}

.social-link {
    color: #3b82f6;
    text-decoration: none;
    font-size: 0.9rem;
    transition: all 0.3s ease;
}

.social-link:hover {
    text-decoration: underline;
}

.social-links a {
    background: linear-gradient(to right, #7f5af0, #3c8ce7);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-weight: 500;
    transition: all 0.3s ease;
}

.social-links a:hover {
    text-shadow: 0 0 5px rgba(123, 31, 162, 0.3);
}

.back-button {
    margin-top: 1rem;
    background: linear-gradient(to right, #7b1fa2, #2196f3);
    color: white;
    border: none;
    border-radius: 20px;
    padding: 0.5rem 1rem;
    cursor: pointer;
    transition: all 0.3s ease;
}

.back-button:hover {
    opacity: 0.9;
    transform: scale(1.02);
}

/* Add this to your CSS section */
.back-button-container {
    display: flex;
    justify-content: center;
    width: 100%;
    margin-top: 1rem;
}




/* 1. Header with gradient and glow */
.header {
    font-size: 2.5rem;
    font-weight: 600;
    background: linear-gradient(to right, #7f5af0, #3c8ce7);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    text-shadow: 0 0 8px rgba(123, 31, 162, 0.4);
    margin-bottom: 1rem;
    text-align: center;
    position: relative;
    animation: header-glow 2s infinite alternate;
}

@keyframes header-glow {
    0% { text-shadow: 0 0 5px rgba(123, 31, 162, 0.4); }
    100% { text-shadow: 0 0 15px rgba(33, 150, 243, 0.6); }
}

/* General styling for cards and containers */
.card, .chat-container, .guide-container, .download-container, .upload-area {
    border-radius: 12px;
    padding: 1rem;
    margin-bottom: 1rem;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
}

/* Gradient border for upload section */
.upload-section {
    border-radius: 15px !important;
    padding: 1.5rem !important;
    margin-bottom: 1.5rem !important;
    background: rgba(30, 30, 30, 0.7) !important;
    border: 1px solid rgba(123, 31, 162, 0.3) !important;
    box-shadow: 0 4px 20px rgba(0, 0, 0, 0.25) !important;
}

.upload-section .stFileUploader > div:first-child {
    display: flex !important;
    justify-content: center !important;
}

/* Style the file uploader dropzone */
.upload-section .stFileUploader > div:last-child {
    text-align: center !important;
    margin-top: 1rem !important;
}
        
.upload-section .stFileUploader button {
    margin: 0 auto !important;
    display: block !important;
    border-radius: 20px !important;
    background: linear-gradient(to right, #7b1fa2, #2196f3) !important;
    color: white !important;
}

/* File list items */
.upload-section .stFileUploader [role='listitem'] {
    background: rgba(40, 40, 40, 0.7) !important;
    border-radius: 8px !important;
    padding: 0.5rem !important;
    margin: 0.25rem 0 !important;
    border: 1px solid rgba(123, 31, 162, 0.2) !important;
}

         
/* Responsive adjustments */
@media (max-width: 768px) {
    .header {
        font-size: 2rem;
    }
}
</style>
""", unsafe_allow_html=True)

def extract_text_from_pdf(pdf_file):
    try:
        doc = fitz.open(stream=pdf_file.getvalue(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        st.error(f"PDF processing error: {str(e)}")
        return ""

def extract_text_from_image(image_file):
    try:
        img = Image.open(image_file)
        text = pytesseract.image_to_string(img)
        return text if text.strip() else "No text could be extracted from the image"
    except Exception as e:
        return f"Image processing error: {str(e)}"

def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(io.BytesIO(docx_file.getvalue()))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"DOCX processing error: {str(e)}"

def extract_text_from_txt(txt_file):
    try:
        return txt_file.getvalue().decode("utf-8")
    except Exception as e:
        return f"Text file processing error: {str(e)}"

def get_pdf_text(uploaded_files):
    text = ""
    for file in uploaded_files:
        if file.name.lower().endswith(('.png', '.jpg', '.jpeg')):
            st.info(f"Processing image: {file.name}")
            file_text = extract_text_from_image(file)
            text += f"Image Content:\n{file_text}\n\n"
        elif file.name.lower().endswith('.pdf'):
            st.info(f"Processing PDF: {file.name}")
            file_text = extract_text_from_pdf(file)
            text += f"PDF Content:\n{file_text}\n\n"
        elif file.name.lower().endswith('.docx'):
            st.info(f"Processing DOCX: {file.name}")
            file_text = extract_text_from_docx(file)
            text += f"DOCX Content:\n{file_text}\n\n"
        elif file.name.lower().endswith('.txt'):
            st.info(f"Processing TXT: {file.name}")
            file_text = extract_text_from_txt(file)
            text += f"TXT Content:\n{file_text}\n\n"
        else:
            st.warning(f"Unsupported file type: {file.name}")
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    return text_splitter.split_text(text)

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=os.getenv("GEMINI_API_KEY")
    )
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")
    st.session_state.vector_store = vector_store
    st.session_state.processed_text = "\n".join(text_chunks)
    return vector_store

def get_conversational_chain():
    prompt_template = """
    You are an expert document assistant named DocuMind. Follow these rules strictly:
    
    1. Provide detailed, accurate answers from the document context
    2. Only include page numbers when explicitly asked by the user
    3. For "what page" or "which page" questions, always provide the page number
    4. Structure answers clearly with bullet points when appropriate
    5. If information isn't found, say "This information was not found in the document"
    
    Context:\n{context}\n
    Question: {question}\n
    
    Response Examples:
    
    User asks: "What is the main conclusion?"
    Response: "The main conclusion is that... (additional details)"
    
    User asks: "What page is the methodology on?"
    Response: "The methodology is described on page 5."
    
    User asks: "Tell me about the results"
    Response: "The results show:
    • Key finding 1 (details)
    • Key finding 2 (details)"
    """
    model = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=0.2,
        google_api_key=os.getenv("GEMINI_API_KEY")
    )
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

def detect_pii(text):
    try:
        results = {
            "emails": re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text),
            "phones": re.findall(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text),
            "credit_cards": re.findall(r'\b(?:\d[ -]*?){13,16}\b', text),
            "names": re.findall(r'\b[A-Z][a-z]+ [A-Z][a-z]+\b', text)
        }
        return results
    except Exception as e:
        st.error(f"PII detection error: {str(e)}")
        return None

def redact_pdf_content(pdf_file):
    try:
        doc = fitz.open(stream=pdf_file.getvalue(), filetype="pdf")
        for page in doc:
            text = page.get_text()
            pii_results = detect_pii(text)
            if pii_results:
                for email in pii_results.get("emails", []):
                    areas = page.search_for(email)
                    [page.add_redact_annot(area, fill=(0, 0, 0)) for area in areas]
                for phone in pii_results.get("phones", []):
                    areas = page.search_for(phone)
                    [page.add_redact_annot(area, fill=(0, 0, 0)) for area in areas]
                page.apply_redactions()
        return doc
    except Exception as e:
        st.error(f"PDF redaction error: {str(e)}")
        return None

def create_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    text = text.encode('latin1', 'replace').decode('latin1')
    pdf.multi_cell(0, 10, text)
    return pdf.output(dest='S').encode('latin1')

def export_content(content, format='txt', filename="export"):
    try:
        if format == 'txt':
            return content.encode('utf-8'), 'text/plain', f"{filename}.txt"
        elif format == 'pdf':
            return create_pdf(content), 'application/pdf', f"{filename}.pdf"
        elif format == 'docx':
            doc = docx.Document()
            doc.add_paragraph(content)
            bio = io.BytesIO()
            doc.save(bio)
            return bio.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f"{filename}.docx"
    except Exception as e:
        st.error(f"Export error: {str(e)}")
        return None, None, None

def reset_all_data():
    try:
        shutil.rmtree("faiss_index", ignore_errors=True)
        st.toast("✅ FAISS index cleared.", icon="✅")
    except Exception as e:
        st.toast(f"Error clearing FAISS index: {str(e)}", icon="❌")
    
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    
    st.toast("Chat history and document data cleared. Please upload new documents to continue.", icon="ℹ️")
    st.rerun()

def clear_chat_history():
    st.session_state.chat_history = []
    st.toast("Chat history cleared. The document analysis remains available for new questions.", icon="ℹ️")
    st.rerun()

def show_project_info():
    st.markdown("""
    <div class="info-card">
        <h3>What is DocuMind?</h3>
        <p>DocuMind is an advanced AI-powered document analysis tool that helps you extract insights from various file formats including PDFs, Word documents, text files, and images. It combines OCR technology with large language models to provide intelligent document understanding, summarization, and question-answering capabilities.</p>
        <p>Key features include:</p>
        <ul>
            <li>Multi-format document processing</li>
            <li>Context-aware question answering</li>
            <li>Automatic PII detection and redaction</li>
            <li>Document summarization</li>
            <li>Conversational interface</li>
        </ul>
    </div>
    
    <div class="info-card">
        <h3>How It Works</h3>
        <p>DocuMind follows a sophisticated pipeline to analyze your documents:</p>
        <ol>
            <li><strong>Document Ingestion</strong>: Upload your files in supported formats</li>
            <li><strong>Text Extraction</strong>: Uses PyMuPDF for PDFs, pytesseract for images, and python-docx for Word documents</li>
            <li><strong>Chunking</strong>: Splits content into manageable sections using RecursiveCharacterTextSplitter</li>
            <li><strong>Vector Embedding</strong>: Creates semantic representations using Google's Gemini embeddings</li>
            <li><strong>Query Processing</strong>: Uses Gemini-Pro model to understand and respond to your questions</li>
            <li><strong>Response Generation</strong>: Provides accurate, context-aware answers with source references</li>
        </ol>
    </div>
    
    <div class="info-card">
        <h3>Technology Stack</h3>
        <p>DocuMind is built with cutting-edge technologies:</p>
        <ul>
            <li><strong>Streamlit</strong>: For the interactive web interface</li>
            <li><strong>Google Gemini</strong>: For embeddings and question answering (models/embedding-001 and gemini-pro)</li>
            <li><strong>LangChain</strong>: For document processing and retrieval pipelines</li>
            <li><strong>FAISS</strong>: For efficient vector similarity search</li>
            <li><strong>PyMuPDF</strong>: For PDF text extraction and redaction</li>
            <li><strong>pytesseract</strong>: For OCR from images</li>
            <li><strong>python-docx</strong>: For Word document processing</li>
        </ul>
        <p>Learn more about these technologies:</p>
        <ul>
            <li><a href="https://streamlit.io/" target="_blank">Streamlit</a></li>
            <li><a href="https://ai.google.dev/" target="_blank">Google Gemini</a></li>
            <li><a href="https://python.langchain.com/" target="_blank">LangChain</a></li>
        </ul>
    </div>
    
   <div class="developer-info">
        <div class="developer-line">
            <span class="developer-name">Developed by Prena Patil</span>
            <a href="https://github.com/prenapatil" target="_blank" class="social-link"> GitHub</a>
            <a href="https://linkedin.com/in/prenapatil" target="_blank" class="social-link">LinkedIn</a>
        </div>
    </div>
    <div class="back-button-container">
    """, unsafe_allow_html=True)

    if st.button("⬅️ Back to Chat", key="back_to_chat"):
        st.session_state.show_project_info = False
        st.rerun()
    st.markdown("</div></div>", unsafe_allow_html=True)
    

def main():
    st.set_page_config("✨ DocuMind AI", page_icon="📄", layout="wide")
    st.markdown('<div class="header">✨DocuMind Ai — Chat with Documents</div>', unsafe_allow_html=True)

    if not st.session_state.get('processed_text') and not st.session_state.get('show_guide') and not st.session_state.get('show_project_info'):
        st.markdown("<h1 style='text-align: center; margin-top: 2rem;'>Welcome to DocuMind AI</h1>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center;'>Upload and process your documents to begin</p>", unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.markdown('<div class="sidebar-header">DocuMind AI</div>', unsafe_allow_html=True)
        st.markdown('<div class="sidebar-subheader">Upload documents and images</div>', unsafe_allow_html=True)
            
        uploaded_files = st.file_uploader(
            "Upload Documents",
            type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            label_visibility="collapsed"
        )
        
        if st.button("Process Documents", use_container_width=True):
            with st.spinner("Processing..."):
                if uploaded_files:
                    st.session_state.original_files = uploaded_files
                    raw_text = get_pdf_text(uploaded_files)
                    if raw_text.strip():
                        text_chunks = get_text_chunks(raw_text)
                        get_vector_store(text_chunks)
                        st.toast("Processing complete!", icon="✅")
                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": "Hello! I've processed your documents. How can I help you today?"
                        })
                        st.rerun()
                    else:
                        st.toast("No text extracted - check file formats", icon="⚠️")
                else:
                    st.toast("Please upload files first", icon="⚠️")
        
        # Divider
        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)
        
        # Chat Management Section
        st.markdown('<div class="sidebar-section-title">Chat Management</div>', unsafe_allow_html=True)
        
        # Download options
        st.markdown('<div class="download-label">Export as:</div>', unsafe_allow_html=True)
        col1, col2 = st.columns([4, 1])
        with col1:
            st.session_state.download_format = st.selectbox(
                "Format",
                ["TXT", "PDF", "DOCX"],
                label_visibility="collapsed"
            )
        with col2:
            export_filename = "document_analysis"
            content = "\n".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.chat_history]) if st.session_state.chat_history else "No chat history yet"
            
            file_data, mime_type, file_name = export_content(
                content,
                st.session_state.download_format.lower(),
                export_filename
            )
            
            if file_data:
                st.download_button(
                    label="⬇️",
                    data=file_data,
                    file_name=file_name,
                    mime=mime_type,
                    use_container_width=True,
                    disabled=not st.session_state.chat_history
                )
        
        if st.button("Clear Chat History", use_container_width=True, type="secondary"):
            clear_chat_history()
        
        if st.button("🔄 Reset All Data", use_container_width=True, type="secondary"):
            reset_all_data()
        
        # Divider
        st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)
        
        # Project Info Section
        st.markdown('<div class="sidebar-section-title">Project Info</div>', unsafe_allow_html=True)
        
        if st.button("About Project", use_container_width=True):
            st.session_state.show_project_info = True
            st.rerun()
    
    # Main content
    if st.session_state.show_project_info:
        show_project_info()
    else:
        # Chat interface
       # Add this import at the top of your file with other imports


# Update your message display code to:
        for message in st.session_state.chat_history:
            if message["role"] == "user":
                st.markdown(f"""
                <div class="user-message-container">
                    <div class="user-message">
                        <div class="message-meta">
                            {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - You
                        </div>
                        <div class="message-content">
                            {message["content"]}
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="bot-message-container">
                    <div class="bot-message">
                        <div class="message-meta">
                            {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - DocuMind
                        </div>
                        <div class="message-content">
                            {message["content"]}
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
        # User input
        user_question = st.chat_input("Ask about your documents...")
        if user_question and st.session_state.vector_store:
            st.session_state.chat_history.append({"role": "user", "content": user_question})
            with st.spinner("Analyzing..."):
                try:
                    if "pii" in user_question.lower() or "hide" in user_question.lower() or "redact" in user_question.lower():
                        if st.session_state.original_files:
                            pdf_file = next((f for f in st.session_state.original_files if f.name.lower().endswith('.pdf')), None)
                            if pdf_file:
                                redacted_doc = redact_pdf_content(pdf_file)
                                if redacted_doc:
                                    pdf_bytes = redacted_doc.write()
                                    b64 = base64.b64encode(pdf_bytes).decode()
                                    pdf_download = f'<a href="data:application/pdf;base64,{b64}" download="redacted_document.pdf">⬇️ Download Redacted PDF</a>'
                                    answer = f"I've redacted PII from the original document. {pdf_download}"
                                else:
                                    answer = "Could not redact the document. Please try again."
                            else:
                                answer = "No PDF file found to redact."
                        else:
                            answer = "No original files available for redaction. Please upload documents again."
                    else:
                        chain = get_conversational_chain()
                        docs = st.session_state.vector_store.similarity_search(user_question)
                        response = chain(
                            {"input_documents": docs, "question": user_question},
                            return_only_outputs=True
                        )
                        answer = response['output_text']
                    
                    st.session_state.chat_history.append({"role": "assistant", "content": answer})
                    st.rerun()
                except Exception as e:
                    st.error(f"Analysis error: {str(e)}")

        # Add this at the bottom of your main() function
                st.markdown("""
                <script>
                // Handle back button click
                function handleBackToChat() {
                    window.parent.postMessage({type: 'streamlit:setComponentValue', value: 'back_to_chat'}, '*');
                }
                </script>
                """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()